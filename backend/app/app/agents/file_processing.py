"""
File processing utilities for knowledge management
"""
from typing import List, Any
import io
import PyPDF2
import docx
import json
import csv
from agno.document import Document


def process_uploaded_file(file_content: bytes, filename: str, content_type: str) -> List[Document]:
    """
    Process uploaded file and convert to documents
    """
    documents = []
    
    try:
        if content_type == "application/pdf" or filename.lower().endswith('.pdf'):
            documents = process_pdf(file_content, filename)
        elif content_type == "text/plain" or filename.lower().endswith('.txt'):
            documents = process_text(file_content, filename)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx'):
            documents = process_docx(file_content, filename)
        elif content_type == "application/json" or filename.lower().endswith('.json'):
            documents = process_json(file_content, filename)
        elif content_type == "text/csv" or filename.lower().endswith('.csv'):
            documents = process_csv(file_content, filename)
        else:
            # Try to process as text
            try:
                text_content = file_content.decode('utf-8')
                documents = [Document(
                    content=text_content,
                    meta_data={
                        "filename": filename,
                        "content_type": content_type,
                        "processing_method": "fallback_text"
                    }
                )]
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file type: {content_type}")
    
    except Exception as e:
        # Create a document with error information
        documents = [Document(
            content=f"Error processing file {filename}: {str(e)}",
            meta_data={
                "filename": filename,
                "content_type": content_type,
                "processing_error": str(e),
                "processing_method": "error"
            }
        )]
    
    return documents


def process_pdf(file_content: bytes, filename: str) -> List[Document]:
    """Process PDF file"""
    documents = []
    
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    content=text,
                    meta_data={
                        "filename": filename,
                        "page_number": page_num + 1,
                        "total_pages": len(pdf_reader.pages),
                        "processing_method": "pdf_extraction"
                    }
                ))
    
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")
    
    return documents


def process_text(file_content: bytes, filename: str) -> List[Document]:
    """Process text file"""
    try:
        text_content = file_content.decode('utf-8')
        
        # Split into chunks if text is very long
        max_chunk_size = 4000  # characters
        chunks = []
        
        if len(text_content) <= max_chunk_size:
            chunks = [text_content]
        else:
            # Split by paragraphs first
            paragraphs = text_content.split('\n\n')
            current_chunk = ""
            
            for paragraph in paragraphs:
                if len(current_chunk + paragraph) <= max_chunk_size:
                    current_chunk += paragraph + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = paragraph + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                content=chunk,
                meta_data={
                    "filename": filename,
                    "chunk_number": i + 1,
                    "total_chunks": len(chunks),
                    "processing_method": "text_chunking"
                }
            ))
        
        return documents
    
    except UnicodeDecodeError:
        raise ValueError("Unable to decode text file as UTF-8")


def process_docx(file_content: bytes, filename: str) -> List[Document]:
    """Process DOCX file"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Combine paragraphs into chunks
        text_content = '\n\n'.join(paragraphs)
        
        return [Document(
            content=text_content,
            meta_data={
                "filename": filename,
                "paragraph_count": len(paragraphs),
                "processing_method": "docx_extraction"
            }
        )]
    
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")


def process_json(file_content: bytes, filename: str) -> List[Document]:
    """Process JSON file"""
    try:
        json_content = json.loads(file_content.decode('utf-8'))
        
        # Convert JSON to readable text
        if isinstance(json_content, dict):
            text_content = json.dumps(json_content, indent=2, ensure_ascii=False)
        elif isinstance(json_content, list):
            # Process each item in the list
            documents = []
            for i, item in enumerate(json_content):
                item_text = json.dumps(item, indent=2, ensure_ascii=False)
                documents.append(Document(
                    content=item_text,
                    meta_data={
                        "filename": filename,
                        "item_index": i,
                        "total_items": len(json_content),
                        "processing_method": "json_item_extraction"
                    }
                ))
            return documents
        else:
            text_content = str(json_content)
        
        return [Document(
            content=text_content,
            meta_data={
                "filename": filename,
                "processing_method": "json_extraction"
            }
        )]
    
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON format: {str(e)}")
    except UnicodeDecodeError:
        raise ValueError("Unable to decode JSON file as UTF-8")


def process_csv(file_content: bytes, filename: str) -> List[Document]:
    """Process CSV file"""
    try:
        csv_content = file_content.decode('utf-8')
        csv_file = io.StringIO(csv_content)
        csv_reader = csv.DictReader(csv_file)
        
        documents = []
        for i, row in enumerate(csv_reader):
            # Convert row to readable text
            row_text = '\n'.join([f"{key}: {value}" for key, value in row.items()])
            
            documents.append(Document(
                content=row_text,
                meta_data={
                    "filename": filename,
                    "row_number": i + 1,
                    "processing_method": "csv_row_extraction",
                    "csv_headers": list(row.keys())
                }
            ))
        
        return documents
    
    except UnicodeDecodeError:
        raise ValueError("Unable to decode CSV file as UTF-8")
    except Exception as e:
        raise ValueError(f"Error processing CSV: {str(e)}")
